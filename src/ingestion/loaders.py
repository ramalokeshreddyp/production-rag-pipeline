from __future__ import annotations

import os
import re
import unicodedata
from pathlib import Path
from typing import List, Optional
from bs4 import BeautifulSoup

from src.core.exceptions import DocumentLoadError
from src.core.types import Document


class TextCleaner:
    """Utilities for sanitizing, cleaning, and normalizing text from various formats."""

    @staticmethod
    def clean(text: str) -> str:
        """Sanitizes raw text: normalizes unicode, handles whitespace, and strips zero-width chars."""
        if not text:
            return ""

        # Normalize unicode (NFKC canonical decomposition & composition)
        text = unicodedata.normalize("NFKC", text)

        # Remove zero-width characters and unusual control characters (keep \n, \t, \r)
        text = re.sub(r"[\u200B-\u200D\uFEFF]", "", text)
        text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)

        # Standardize line breaks
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Reduce excessive blank lines (more than 2 consecutive newlines -> 2)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Normalize intra-line spaces (multiple tabs/spaces -> single space, preserving newlines)
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
        return "\n".join(lines).strip()

    @staticmethod
    def strip_html(html_content: str) -> str:
        """Extracts clean text from HTML, removing script and style tags."""
        soup = BeautifulSoup(html_content, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
            tag.decompose()
        return soup.get_text(separator="\n")


class DocumentLoader:
    """Multi-format document loader supporting TXT, MD, PDF, HTML, and DOCX."""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".html", ".htm", ".docx"}

    def __init__(self, clean_text: bool = True):
        self.clean_text = clean_text

    def load_file(self, file_path: str | Path) -> Document:
        """Loads and extracts text content from a single file."""
        path = Path(file_path)
        if not path.exists():
            raise DocumentLoadError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise DocumentLoadError(
                f"Unsupported file format '{ext}'. Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        try:
            if ext in {".txt", ".md"}:
                raw_text = self._load_text_file(path)
            elif ext in {".html", ".htm"}:
                raw_text = self._load_html_file(path)
            elif ext == ".pdf":
                raw_text = self._load_pdf_file(path)
            elif ext == ".docx":
                raw_text = self._load_docx_file(path)
            else:
                raw_text = self._load_text_file(path)

            content = TextCleaner.clean(raw_text) if self.clean_text else raw_text

            file_stat = path.stat()
            metadata = {
                "source": str(path.resolve()),
                "filename": path.name,
                "file_type": ext,
                "file_size_bytes": file_stat.st_size,
                "modified_time": file_stat.st_mtime,
            }

            return Document(
                content=content,
                metadata=metadata,
                source_path=str(path.resolve()),
                filename=path.name,
                file_type=ext,
            )

        except Exception as e:
            if isinstance(e, DocumentLoadError):
                raise e
            raise DocumentLoadError(f"Failed to load document '{path}': {str(e)}") from e

    def load_directory(
        self,
        directory_path: str | Path,
        recursive: bool = True,
        extensions: Optional[List[str]] = None,
    ) -> List[Document]:
        """Loads all supported documents from a directory."""
        dir_path = Path(directory_path)
        if not dir_path.exists() or not dir_path.is_dir():
            raise DocumentLoadError(f"Directory not found or invalid: {dir_path}")

        target_extensions = (
            {ext.lower() if ext.startswith(".") else f".{ext.lower()}" for ext in extensions}
            if extensions
            else self.SUPPORTED_EXTENSIONS
        )

        documents: List[Document] = []
        pattern = "**/*" if recursive else "*"
        for path in dir_path.glob(pattern):
            if path.is_file() and path.suffix.lower() in target_extensions:
                try:
                    doc = self.load_file(path)
                    if doc.content.strip():
                        documents.append(doc)
                except Exception as exc:
                    # Log or skip corrupt files gracefully
                    print(f"Warning: Skipped loading {path} due to error: {exc}")

        return documents

    def _load_text_file(self, path: Path) -> str:
        """Reads plain text / markdown file with encoding fallback."""
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentLoadError(f"Unable to decode text file {path} with common encodings.")

    def _load_html_file(self, path: Path) -> str:
        """Reads and strips HTML content."""
        html_raw = self._load_text_file(path)
        return TextCleaner.strip_html(html_raw)

    def _load_pdf_file(self, path: Path) -> str:
        """Extracts text from PDF using pypdf."""
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            return "\n\n".join(pages_text)
        except ImportError:
            raise DocumentLoadError("pypdf is required to load PDF documents. Install via requirements.txt.")
        except Exception as e:
            raise DocumentLoadError(f"Failed to parse PDF {path}: {str(e)}")

    def _load_docx_file(self, path: Path) -> str:
        """Extracts text from DOCX file."""
        try:
            import docx
            doc = docx.Document(str(path))
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            raise DocumentLoadError("python-docx is required to load DOCX files. Install via requirements.txt.")
        except Exception as e:
            raise DocumentLoadError(f"Failed to parse DOCX {path}: {str(e)}")
